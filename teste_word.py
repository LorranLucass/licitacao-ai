from docx import Document

doc = Document("interface/modelos/proposta.docx")

print("\nPARÁGRAFOS:")
print("=" * 50)

for i, p in enumerate(doc.paragraphs):
    texto = p.text.strip()

    if texto:
        print(f"[{i}] {texto}")

print("\nTABELAS:")
print("=" * 50)

for t, tabela in enumerate(doc.tables):

    print(f"\nTABELA {t}")

    for l, linha in enumerate(tabela.rows):

        valores = []

        for c, celula in enumerate(linha.cells):
            valores.append(celula.text.strip())

        print(f"Linha {l}: {valores}")