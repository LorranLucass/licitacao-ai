from pathlib import Path

from docx import Document

from services.bloco_notas import (
    formatar_data_hora_sessao,
    calcular_valor_total_proposta,
    formatar_moeda,
    converter_numero,
    converter_quantidade,
)


def calcular_preco_item(item):
    """
    Calcula o preço unitário e total de UM item, com a
    mesma regra usada no bloco de notas e no valor total da
    proposta (para nunca mostrar números diferentes em
    lugares diferentes):

    - se o edital trouxe valor estimado, usa esse valor;
    - senão, usa custo × 1,60 (markup padrão).
    """

    quantidade = converter_quantidade(
        item.get("quantidade", 0)
    )

    custo = converter_numero(
        item.get("custo", 0)
    )

    valor_unitario = item.get("valor_unitario")
    valor_total = item.get("valor_total")

    tem_estimado = (
        valor_unitario is not None
        and str(valor_unitario).strip() != ""
    )

    if tem_estimado:

        preco_unitario = converter_numero(valor_unitario)

        if valor_total is not None and str(valor_total).strip():
            preco_total = converter_numero(valor_total)
        else:
            preco_total = preco_unitario * quantidade

    else:

        preco_unitario = custo * 1.60
        preco_total = preco_unitario * quantidade

    return preco_unitario, preco_total


# ============================================================
# NÚMERO POR EXTENSO (REAIS)
# ============================================================

_UNIDADES = [
    "", "um", "dois", "três", "quatro", "cinco",
    "seis", "sete", "oito", "nove"
]

_DEZ_A_DEZENOVE = [
    "dez", "onze", "doze", "treze", "catorze", "quinze",
    "dezesseis", "dezessete", "dezoito", "dezenove"
]

_DEZENAS = [
    "", "", "vinte", "trinta", "quarenta", "cinquenta",
    "sessenta", "setenta", "oitenta", "noventa"
]

_CENTENAS = [
    "", "cento", "duzentos", "trezentos", "quatrocentos",
    "quinhentos", "seiscentos", "setecentos", "oitocentos",
    "novecentos"
]


def _grupo_extenso(numero):
    """Converte um número de 0 a 999 por extenso."""

    if numero == 0:
        return ""

    if numero == 100:
        return "cem"

    centena = numero // 100
    resto = numero % 100

    partes = []

    if centena:
        partes.append(_CENTENAS[centena])

    if resto:

        if resto < 10:
            partes.append(_UNIDADES[resto])

        elif resto < 20:
            partes.append(_DEZ_A_DEZENOVE[resto - 10])

        else:

            dezena = resto // 10
            unidade = resto % 10

            if unidade:
                partes.append(
                    _DEZENAS[dezena] + " e " + _UNIDADES[unidade]
                )
            else:
                partes.append(_DEZENAS[dezena])

    return " e ".join(partes)


def numero_por_extenso(numero):
    """
    Converte um número inteiro (0 a 999.999.999) por extenso,
    em português.
    """

    numero = int(numero)

    if numero == 0:
        return "zero"

    milhoes = numero // 1_000_000
    milhares = (numero % 1_000_000) // 1000
    resto = numero % 1000

    partes = []

    if milhoes:

        texto = _grupo_extenso(milhoes)

        rotulo_milhao = (
            "milhão" if milhoes == 1 else "milhões"
        )

        # "de" só entra quando o milhão é o último grupo
        # (ex: "um milhão de reais"), e não quando há
        # milhares/unidades depois (ex: "um milhão e
        # duzentos mil reais").
        if milhares == 0 and resto == 0:
            partes.append(f"{texto} {rotulo_milhao} de")
        else:
            partes.append(f"{texto} {rotulo_milhao}")

    if milhares:

        texto = _grupo_extenso(milhares)

        partes.append(
            "mil" if milhares == 1 else f"{texto} mil"
        )

    if resto:
        partes.append(_grupo_extenso(resto))

    return " ".join(partes) if partes else "zero"


def valor_por_extenso(valor):
    """
    Converte um valor em reais por extenso.

    Exemplo: 250499.00 ->
    "DUZENTOS E CINQUENTA MIL QUATROCENTOS E NOVENTA E NOVE REAIS"
    """

    valor = round(float(valor or 0), 2)

    reais = int(valor)
    centavos = round((valor - reais) * 100)

    texto_reais = numero_por_extenso(reais)
    sufixo_reais = "real" if reais == 1 else "reais"

    resultado = f"{texto_reais} {sufixo_reais}"

    if centavos:

        texto_centavos = numero_por_extenso(centavos)
        sufixo_centavos = (
            "centavo" if centavos == 1 else "centavos"
        )

        resultado += (
            f" e {texto_centavos} {sufixo_centavos}"
        )

    return resultado.upper()


def substituir_texto_no_paragrafo(paragrafo, substituicoes):
    """
    Substitui marcadores no parágrafo preservando a
    formatação de cada run (negrito, fonte, tamanho etc.).

    O Word costuma dividir um mesmo marcador em vários runs
    internos quando o documento foi editado manualmente
    (ex: "{{ORGAO}}" pode virar 3 runs diferentes: "{{ORG",
    "AO", "}}"). Uma troca simples via parágrafo.text falha
    nesse caso e destrói a formatação. Esta função localiza
    o marcador no texto completo do parágrafo, descobre
    quais runs ele atravessa, e edita só esses runs.
    """

    for marcador, valor in substituicoes.items():

        valor = "" if valor is None else str(valor)

        # Repete até não haver mais ocorrências desse
        # marcador neste parágrafo (pode aparecer mais de
        # uma vez).
        while True:

            runs = paragrafo.runs

            texto_completo = "".join(
                run.text for run in runs
            )

            posicao = texto_completo.find(marcador)

            if posicao == -1:
                break

            fim = posicao + len(marcador)

            # Mapeia cada run para o intervalo [inicio, fim)
            # que ele ocupa dentro do texto completo do
            # parágrafo, e guarda os que o marcador atravessa.
            indice = 0
            runs_afetados = []

            for run in runs:

                inicio_run = indice
                fim_run = indice + len(run.text)

                if fim_run > posicao and inicio_run < fim:
                    runs_afetados.append(
                        (run, inicio_run, fim_run)
                    )

                indice = fim_run

            if not runs_afetados:
                # Segurança: não deveria acontecer, já que
                # achamos o marcador no texto completo.
                break

            for numero, (run, inicio_run, fim_run) in enumerate(
                runs_afetados
            ):

                # Texto do run que fica ANTES do marcador.
                antes = run.text[
                    : max(0, posicao - inicio_run)
                ]

                # Texto do run que fica DEPOIS do marcador.
                depois = run.text[
                    max(0, fim - inicio_run):
                ]

                if numero == 0:

                    # O valor novo entra inteiro no primeiro
                    # run afetado — assim ele herda a
                    # formatação desse run.
                    run.text = antes + valor + depois

                else:

                    # Os demais runs só mantêm o que sobrou
                    # fora do marcador.
                    run.text = antes + depois


def substituir_marcadores(documento, substituicoes):
    """
    Substitui marcadores em:
    - parágrafos do corpo do documento
    - tabelas do corpo do documento
    - cabeçalho e rodapé de cada seção (incluindo tabelas
      dentro deles, se houver)
    """

    # --------------------------------------------------------
    # PARÁGRAFOS DO CORPO
    # --------------------------------------------------------

    for paragrafo in documento.paragraphs:

        substituir_texto_no_paragrafo(
            paragrafo,
            substituicoes
        )

    # --------------------------------------------------------
    # TABELAS DO CORPO
    # --------------------------------------------------------

    for tabela in documento.tables:

        for linha in tabela.rows:

            for celula in linha.cells:

                for paragrafo in celula.paragraphs:

                    substituir_texto_no_paragrafo(
                        paragrafo,
                        substituicoes
                    )

    # --------------------------------------------------------
    # CABEÇALHO E RODAPÉ
    # --------------------------------------------------------

    for secao in documento.sections:

        for area in [secao.header, secao.footer]:

            for paragrafo in area.paragraphs:

                substituir_texto_no_paragrafo(
                    paragrafo,
                    substituicoes
                )

            for tabela in area.tables:

                for linha in tabela.rows:

                    for celula in linha.cells:

                        for paragrafo in celula.paragraphs:

                            substituir_texto_no_paragrafo(
                                paragrafo,
                                substituicoes
                            )


def localizar_tabela_itens(documento):
    """
    Procura a tabela de itens existente no modelo.

    Não cria uma tabela nova.

    Procura uma tabela que possua:
    ITEM
    DESCRIÇÃO
    """

    for tabela in documento.tables:

        for linha in tabela.rows[:5]:

            texto = " ".join(
                celula.text.upper()
                for celula in linha.cells
            )

            if (
                "ITEM" in texto
                and "DESCRI" in texto
            ):
                return tabela

    return None


def localizar_linha_cabecalho(tabela):
    """
    Localiza a linha que contém o cabeçalho
    da tabela de itens.
    """

    for indice, linha in enumerate(tabela.rows):

        texto = " ".join(
            celula.text.upper()
            for celula in linha.cells
        )

        if (
            "ITEM" in texto
            and "DESCRI" in texto
        ):
            return indice

    return None


def localizar_colunas(tabela, linha_cabecalho):
    """
    Identifica as colunas da tabela.
    """

    colunas = {}

    linha = tabela.rows[linha_cabecalho]

    for indice, celula in enumerate(linha.cells):

        texto = (
            celula.text
            .strip()
            .upper()
            .replace("\n", " ")
        )

        if "ITEM" in texto:

            colunas["item"] = indice

        elif "DESCRI" in texto:

            colunas["descricao"] = indice

        elif (
            "MARCA" in texto
            or "MODELO" in texto
        ):

            colunas["marca_modelo"] = indice

        elif (
            "QTD" in texto
            or "QUANT" in texto
        ):

            colunas["quantidade"] = indice

        elif (
            "UND" in texto
            or "UNID" in texto
        ):

            colunas["unidade"] = indice

        elif "UNIT" in texto:

            colunas["preco_unitario"] = indice

        elif "TOTAL" in texto:

            colunas["preco_total"] = indice

    return colunas


def limpar_linhas_de_itens(tabela, linha_cabecalho):
    """
    Remove somente as linhas de itens existentes
    abaixo do cabeçalho.

    Mantém a linha TOTAL GERAL.
    """

    linhas_para_remover = []

    for indice in range(
        linha_cabecalho + 1,
        len(tabela.rows)
    ):

        linha = tabela.rows[indice]

        texto = " ".join(
            celula.text.upper()
            for celula in linha.cells
        )

        if "TOTAL GERAL" in texto:
            continue

        linhas_para_remover.append(
            linha
        )

    for linha in linhas_para_remover:

        tabela._tbl.remove(
            linha._tr
        )


def preencher_tabela_itens(documento, itens):
    """
    Preenche a tabela de itens EXISTENTE no modelo.

    Não cria uma tabela nova.

    A estrutura esperada é:

    ITEM
    DESCRIÇÃO
    MARCA / MODELO
    QTD
    UND
    PREÇO UNITÁRIO
    PREÇO TOTAL
    """

    if not itens:
        print(
            "Nenhum item para colocar na proposta."
        )
        return

    tabela = localizar_tabela_itens(
        documento
    )

    if tabela is None:

        print(
            "Tabela de itens não encontrada "
            "no modelo Word."
        )

        return

    linha_cabecalho = localizar_linha_cabecalho(
        tabela
    )

    if linha_cabecalho is None:

        print(
            "Cabeçalho da tabela não encontrado."
        )

        return

    colunas = localizar_colunas(
        tabela,
        linha_cabecalho
    )

    print()
    print(
        "Tabela de itens encontrada."
    )

    print(
        "Colunas:",
        colunas
    )

    # --------------------------------------------------------
    # IDENTIFICAR TOTAL GERAL
    # --------------------------------------------------------

    linha_total = None

    for linha in tabela.rows:

        texto = " ".join(
            celula.text.upper()
            for celula in linha.cells
        )

        if "TOTAL GERAL" in texto:

            linha_total = linha
            break

    # --------------------------------------------------------
    # REMOVER ITENS ANTIGOS
    # --------------------------------------------------------

    limpar_linhas_de_itens(
        tabela,
        linha_cabecalho
    )

    # --------------------------------------------------------
    # INSERIR OS NOVOS ITENS
    # --------------------------------------------------------

    for item in itens:

        nova_linha = tabela.add_row()

        descricao = (
            item.get("descricao")
            or item.get("produto_edital")
            or item.get("produto")
            or ""
        )

        produto = (
            item.get("produto_tabela")
            or ""
        )

        marca = (
            item.get("marca_tabela")
            or ""
        )

        modelo = (
            item.get("modelo")
            or ""
        )

        marca_modelo = marca

        if modelo:

            if marca_modelo:

                marca_modelo += (
                    f" / {modelo}"
                )

            else:

                marca_modelo = modelo

        quantidade = (
            item.get("quantidade")
            or item.get("qtd")
            or ""
        )

        unidade = (
            item.get("unidade")
            or item.get("und")
            or "UND"
        )

        preco_unitario_num, preco_total_num = calcular_preco_item(
            item
        )

        preco_unitario = formatar_moeda(preco_unitario_num)

        preco_total = formatar_moeda(preco_total_num)

        valores = {

            "item": item.get(
                "item",
                ""
            ),

            "descricao": descricao,

            "marca_modelo": marca_modelo,

            "quantidade": quantidade,

            "unidade": unidade,

            "preco_unitario": preco_unitario,

            "preco_total": preco_total,
        }

        for campo, coluna in colunas.items():

            if coluna >= len(
                nova_linha.cells
            ):
                continue

            nova_linha.cells[
                coluna
            ].text = str(
                valores.get(
                    campo,
                    ""
                )
            )

    print(
        f"{len(itens)} item(ns) inserido(s) "
        "na tabela da proposta."
    )

    # --------------------------------------------------------
    # REPOSICIONAR "TOTAL GERAL" NO FIM DA TABELA
    # --------------------------------------------------------
    # tabela.add_row() sempre acrescenta a nova linha no fim
    # da tabela — então, depois de inserir os itens, a linha
    # TOTAL GERAL (que ficou entre o cabeçalho e os itens
    # novos) precisa ser movida de volta pro final.

    if linha_total is not None:

        tabela._tbl.remove(
            linha_total._tr
        )

        tabela._tbl.append(
            linha_total._tr
        )

        # ----------------------------------------------------
        # ATUALIZAR O VALOR DO TOTAL GERAL
        # ----------------------------------------------------
        # O modelo trazia o total do exemplo antigo
        # (R$ 250.499,00) fixo nessas células. Substitui pelo
        # valor recalculado com os itens desta análise.

        valor_total_geral = calcular_valor_total_proposta(
            itens
        )

        texto_total_geral = formatar_moeda(
            valor_total_geral
        )

        for campo in [
            "quantidade",
            "unidade",
            "preco_unitario",
            "preco_total",
        ]:

            coluna = colunas.get(campo)

            if (
                coluna is not None
                and coluna < len(linha_total.cells)
            ):

                linha_total.cells[
                    coluna
                ].text = texto_total_geral


def gerar_proposta_word(
    dados,
    caminho_template,
    caminho_saida,
):
    """
    Gera a proposta Word usando o modelo existente.
    """

    caminho_template = Path(
        caminho_template
    )

    caminho_saida = Path(
        caminho_saida
    )

    print()
    print(
        "================================"
    )
    print(
        "GERANDO PROPOSTA WORD"
    )
    print(
        "================================"
    )

    print(
        "Modelo:",
        caminho_template
    )

    print(
        "Saída:",
        caminho_saida
    )

    if not caminho_template.exists():

        raise FileNotFoundError(
            "Modelo Word não encontrado: "
            f"{caminho_template}"
        )

    caminho_saida.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    documento = Document(
        str(caminho_template)
    )

    # --------------------------------------------------------
    # DADOS
    # --------------------------------------------------------

    orgao = (
        dados.get("orgao")
        or dados.get("órgão")
        or ""
    )

    pregao = (
        dados.get("pregao")
        or dados.get("pregão")
        or dados.get("pregao_dispensa")
        or ""
    )

    objeto = (
        dados.get("objeto")
        or ""
    )

    validade_numero = (
        dados.get("proposta_validade")
        or dados.get("validade_proposta")
        or dados.get("prazo_validade")
        or dados.get("validade")
        or ""
    )

    validade_unidade = (
        dados.get("proposta_validade_unidade")
        or "DIAS"
    )

    validade = (
        f"{validade_numero} {validade_unidade}"
        if validade_numero
        else ""
    )

    entrega = (
        dados.get("entrega_fornecimento")
        or dados.get("prazo_entrega")
        or dados.get("entrega")
        or ""
    )

    pagamento = (
        dados.get("pagamento")
        or dados.get("prazo_pagamento")
        or ""
    )

    garantia_info = dados.get("atencao", {})

    if not isinstance(garantia_info, dict):
        garantia_info = {}

    garantia_situacao = (
        garantia_info.get("garantia")
        or dados.get("garantia")
        or dados.get("prazo_garantia")
        or ""
    )

    garantia_tipos = garantia_info.get("garantia_tipos", [])

    if not isinstance(garantia_tipos, list):
        garantia_tipos = []

    if garantia_tipos:
        garantia = (
            f"{garantia_situacao} - "
            f"{', '.join(garantia_tipos)}"
        )
    else:
        garantia = garantia_situacao

    cidade = (
        dados.get("cidade_estado")
        or dados.get("cidade")
        or dados.get("municipio")
        or "BRASÍLIA/DF"
    )

    horario = (
        dados.get("horario_sessao")
        or ""
    )

    data = (
        dados.get("data_proposta")
        or formatar_data_hora_sessao(dados)
        or dados.get("data")
        or ""
    )

    itens = dados.get("itens", [])

    valor_total = calcular_valor_total_proposta(itens)
    valor_total_texto = formatar_moeda(valor_total)
    valor_total_extenso = valor_por_extenso(valor_total)

    # --------------------------------------------------------
    # MARCADORES
    # --------------------------------------------------------

    substituicoes = {

        "{{ORGAO}}": orgao,

        "{{ÓRGAO}}": orgao,

        "{{PREGAO}}": pregao,

        "{{PREGÃO}}": pregao,

        "{{PREGAO_DISPENSA}}": pregao,

        "{{PREGÃO_DISPENSA}}": pregao,

        "{{OBJETO}}": objeto,

        "{{VALIDADE}}": validade,

        "{{PROPOSTA_VALIDADE}}": validade,

        "{{ENTREGA}}": entrega,

        "{{ENTREGA_FORNECIMENTO}}": entrega,

        "{{PAGAMENTO}}": pagamento,

        "{{GARANTIA}}": garantia,

        "{{DATA}}": data,

        "{{HORARIO}}": horario,

        "{{HORÁRIO}}": horario,

        "{{VALOR_TOTAL}}": valor_total_texto,

        "{{VALOR_TOTAL_EXTENSO}}": valor_total_extenso,

        "{{CIDADE}}": cidade,

    }

    # --------------------------------------------------------
    # SUBSTITUIR MARCADORES
    # --------------------------------------------------------

    substituir_marcadores(
        documento,
        substituicoes
    )

    # --------------------------------------------------------
    # ITENS
    # --------------------------------------------------------

    itens = dados.get(
        "itens",
        []
    )

    preencher_tabela_itens(
        documento,
        itens
    )

    # --------------------------------------------------------
    # SALVAR
    # --------------------------------------------------------

    documento.save(
        str(caminho_saida)
    )

    if not caminho_saida.exists():

        raise RuntimeError(
            "O Word foi processado, "
            "mas o arquivo de saída não foi criado."
        )

    print()
    print(
        "PROPOSTA WORD CRIADA:"
    )

    print(
        caminho_saida
    )

    print()

    return str(
        caminho_saida
    )